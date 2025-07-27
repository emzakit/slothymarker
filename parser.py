import os
import re
import docx
import fitz
import html
from dataclasses import dataclass

@dataclass(unsafe_hash=True)
class Highlight:
    text: str
    start_pos: int = -1
    start_time: float = -1.0
    end_time: float = -1.0
    display_text: str = ""
    sort_key: int = 0

    def __post_init__(self):
        if not self.display_text: self.display_text = self.text

def _create_styled_span(text, index, is_selected, highlight_color, selection_color):
    bg_color = selection_color if is_selected else highlight_color
    style = f"background-color:{bg_color};"
    if is_selected:
        style += " color:white;"
    
    escaped_text = html.escape(text).replace('\n', '<br>')
    return f'<a href="slothy:highlight_{index}" style="color:inherit; text-decoration:none;"><span style="{style}">{escaped_text}</span></a>'

def _parse_simple(content: str) -> tuple[str, list[Highlight]]:
    raw_text = re.sub(r'==(.*?)==', r'\1', flags=re.DOTALL, string=content)
    highlight_texts = re.findall(r'==(.*?)==', content, re.DOTALL)
    highlights = [Highlight(text=h, start_pos=raw_text.find(h)) for h in highlight_texts]
    return raw_text, highlights

def _docx_parser(filepath: str) -> tuple[str, list[str]]:
    doc = docx.Document(filepath)
    raw_text = "\n\n".join([para.text for para in doc.paragraphs])
    highlights_list = [run.text for para in doc.paragraphs for run in para.runs if run.font.highlight_color]
    return raw_text, [h.strip() for h in highlights_list if h.strip()]

def _pdf_parser(filepath: str) -> tuple[str, list[str]]:
    doc = fitz.open(filepath)
    full_text = ""
    highlights_list = []
    for page in doc:
        full_text += page.get_text("text") + "\n"
        for annot in page.annots(types=[fitz.ANNOT_HIGHLIGHT]):
            text = " ".join(page.get_text("text", clip=quad).strip() for quad in annot.quads() or [])
            if text: highlights_list.append(text)
    return full_text, [h.strip() for h in highlights_list if h.strip()]

def render_document_with_highlights(raw_text: str, all_highlights: list[Highlight], selected_highlights: list[Highlight], highlight_color: str, selection_color: str) -> str:
    rendered_text = html.escape(raw_text)
    sorted_by_len = sorted(all_highlights, key=lambda h: len(h.text), reverse=True)
    
    for h_obj in sorted_by_len:
        is_selected = (h_obj in selected_highlights)
        try:
            stable_index = all_highlights.index(h_obj)
            replacement_html = _create_styled_span(h_obj.text, stable_index, is_selected, highlight_color, selection_color)
            escaped_search_text = html.escape(h_obj.text)
            rendered_text = re.sub(re.escape(escaped_search_text), replacement_html, rendered_text, 1)
        except (ValueError, re.error):
            continue
            
    return rendered_text.replace('\n', '<br>')

def parse_document(filepath: str, file_tags: list) -> tuple[str, list, str]:
    from transcript_parser import parse_transcript_file
    extension = os.path.splitext(filepath)[1].lower()
    
    if extension in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        # --- NEW: Strip the invisible header before parsing ---
        header_pattern = re.compile(r'<!---.*?-->\n*', re.DOTALL)
        content = header_pattern.sub('', content)
        # ---
        
        first_line = content.lstrip().split('\n', 1)[0].strip()

        if first_line in file_tags:
            raw_text = re.sub(r'==(.*?)==', r'\1', flags=re.DOTALL, string=content)
            highlights = parse_transcript_file(content)
            return raw_text, highlights, first_line
        else:
            raw_text, highlights = _parse_simple(content)
            return raw_text, highlights, "simple"
            
    elif extension == '.docx':
        raw_text, highlights_text = _docx_parser(filepath)
        highlights = [Highlight(text=h, start_pos=raw_text.find(h)) for h in highlights_text]
        return raw_text, highlights, "simple"

    elif extension == '.pdf':
        raw_text, highlights_text = _pdf_parser(filepath)
        highlights = [Highlight(text=h, start_pos=raw_text.find(h)) for h in highlights_text]
        return raw_text, highlights, "simple"
        
    raise ValueError(f"Unsupported file type: '{extension}'.")